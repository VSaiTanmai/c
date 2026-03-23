"""
CLIF Report Generator — Sections 1-5 (~2 pages each, 6-criteria aligned)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = os.path.join(os.path.dirname(__file__), "CLIF_Report_v3.docx")
doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────
s = doc.styles["Normal"]
s.font.name = "Calibri"; s.font.size = Pt(11)
s.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.15
for lv in range(1,4):
    h = doc.styles[f"Heading {lv}"]
    h.font.color.rgb = RGBColor(0x0B,0x2D,0x5B); h.font.name = "Calibri"
    h.font.bold = True; h.font.size = Pt([0,18,14,12][lv])

def P(text, bold=False, italic=False, sa=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p

def RP(segs, sa=None):
    """segs: [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for t,b,i in segs:
        r = p.add_run(t); r.bold = b; r.italic = i
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p

def B(text):
    return doc.add_paragraph(text, style="List Bullet")

def T(hdrs, rows):
    t = doc.add_table(rows=1, cols=len(hdrs))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for rd in rows:
        row = t.add_row()
        for i,v in enumerate(rd):
            row.cells[i].text = str(v)
            for p in row.cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CLIF"); r.bold=True; r.font.size=Pt(36); r.font.color.rgb=RGBColor(0x0B,0x2D,0x5B)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cognitive Log Investigation Framework"); r.font.size=Pt(20); r.font.color.rgb=RGBColor(0x2E,0x74,0xB5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Powered Multi-Agent SIEM for Autonomous\nThreat Detection, Investigation, and Response")
r.font.size=Pt(13); r.italic=True
for _ in range(3): doc.add_paragraph()
for l,v in [
    ("Problem Statement","SIH1733 — AI-Based Log Investigation Framework for Next-Gen Cyber Forensics"),
    ("Ministry","Ministry of Home Affairs / Bureau of Police Research & Development"),
    ("Event","Smart India Hackathon 2024 — Grand Finale"),
    ("Institution","[Your Institution Name]"),
    ("Team Name","[Your Team Name]"),
    ("Team Leader","[Team Leader Name]"),
    ("Mentor","[Mentor Name]"),
    ("Date","February 2026"),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{l}: "); r.bold=True; r.font.size=Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# ── TOC ───────────────────────────────────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
for x in [
    "1. Project Overview and Team Details",
    "2. Problem Statement and Background",
    "3. Literature Review / Existing Solutions",
    "4. Proposed Solution and Technical Architecture",
    "5. Innovation and Novelty Elements",
    "6. USP vis-à-vis Existing Solutions & Industry Relevance",
    "7. Prototype Demonstration and Real-World Deployment",
    "8. Limitations and Challenges",
    "9. Roadmap Towards MVP",
]:
    p = doc.add_paragraph(x); p.paragraph_format.space_after = Pt(3)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1: PROJECT OVERVIEW AND TEAM DETAILS  (~2 pages)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Project Overview and Team Details", level=1)

doc.add_heading("1.1 Project Summary", level=2)
P(
    "CLIF (Cognitive Log Investigation Framework) is a production-grade, AI-powered SIEM platform "
    "that ingests logs from diverse sources at 56,000+ events per second, enables AI-driven semantic "
    "search over 494,000+ vectorized log entries, and provides cryptographic evidence anchoring via "
    "SHA-256 Merkle trees. Built for SIH1733 under the Ministry of Home Affairs / BPR&D, CLIF bridges "
    "the gap between expensive enterprise SIEMs ($50K–$500K/yr) and capability-limited open-source "
    "alternatives."
)
P(
    "The platform is organized into three architectural planes. The Data Plane handles collection "
    "(Vector aggregator, Tetragon eBPF), streaming (3-broker Redpanda cluster), ingestion (3 Python "
    "consumers), storage (2-node ClickHouse cluster), cold archival (MinIO S3), and evidence anchoring "
    "(Merkle service). The Intelligence Plane houses a four-agent AI pipeline — Triage, Hunter, "
    "Verifier, Reporter — orchestrated via DSPy with LanceDB for RAG-based semantic retrieval. The "
    "Presentation Plane provides a 14-page Next.js 14 SOC dashboard with live event streaming, "
    "semantic search, alert management, attack graph visualization, and system health monitoring. "
    "All 18 services are containerized via Docker Compose with Prometheus/Grafana observability."
)

doc.add_heading("1.2 Core Capabilities", level=2)
for title, desc in [
    ("Multi-Source Ingestion: ",
     "Seven source types (syslog, HTTP, Docker, file, journald, kernel eBPF via Tetragon) normalized "
     "into the CLIF Common Schema via Vector VRL transforms."),
    ("High-Throughput Pipeline: ",
     "3-broker Redpanda (4 topics, 12 partitions, RF=3) with 3 Python consumers — 56,612 EPS "
     "sustained, zero data loss across 2.5M events."),
    ("Columnar Analytics: ",
     "ClickHouse 24.8 (2-node replicated shard), sub-100ms queries, 15-20x ZSTD compression, "
     "TTL tiering: hot 7d → warm 30d → cold S3 90d+."),
    ("AI Semantic Search: ",
     "LanceDB indexes 494K+ entries as 384-dim vectors (all-MiniLM-L6-v2) for natural language "
     "queries like \"suspicious lateral movement after midnight\"."),
    ("Tamper-Proof Evidence: ",
     "SHA-256 Merkle trees over log batches, stored in MinIO with S3 Object Lock (WORM) — any "
     "modification mathematically detectable."),
    ("Multi-Agent AI: ",
     "Triage → Hunter → Verifier → Reporter pipeline via DSPy. Each agent provides explainable "
     "confidence scoring and structured investigation rationale."),
    ("SOC Dashboard: ",
     "14-page Next.js 14 interface — live feed, semantic search, alerts, attack graph (React Flow), "
     "evidence chain, system health, CSV/JSON export."),
]:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(desc)

doc.add_heading("1.3 Team", level=2)
P("[Fill in team details.]", italic=True)
T(["Role","Name","Responsibilities"],[
    ("Team Leader","[Name]","Architecture, design decisions, coordination"),
    ("Backend Developer","[Name]","Consumer pipeline, ClickHouse, Redpanda"),
    ("AI/ML Engineer","[Name]","LanceDB, semantic search, agent pipeline"),
    ("Frontend Developer","[Name]","Next.js dashboard, API routes, visualizations"),
    ("DevOps Engineer","[Name]","Docker orchestration, Prometheus/Grafana"),
    ("Security Researcher","[Name]","MITRE mapping, Merkle evidence, threat modeling"),
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2: PROBLEM STATEMENT AND BACKGROUND  (~2 pages)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Problem Statement and Background", level=1)

doc.add_heading("2.1 Problem Statement (SIH1733)", level=2)
RP([("Title: ",True,False),
    ("Artificial Intelligence Based Log Investigation Framework for Next-Generation Cyber Forensics",False,True)])
P(
    "With the exponential growth of cybercrimes and the proliferation of devices, massive amounts of logs are "
    "generated from computers, mobile devices, IoT, servers, and applications. These logs are crucial for "
    "reconstructing timelines, understanding the nature of incidents, and drawing actionable insights. Current "
    "manual or semi-automated methods are time-consuming and inefficient. There is a pressing need for an "
    "AI-powered framework that can ingest, parse, and analyze logs, enabling investigators to correlate events, "
    "detect anomalies, and generate actionable forensic insights efficiently."
)

doc.add_heading("2.2 Background", level=2)
P(
    "India's cybersecurity landscape faces escalating pressure — CERT-In reported 13.9 lakh incidents in "
    "2022, with government and critical infrastructure as prime targets. SOC analysts are overwhelmed: "
    "enterprises generate 10,000–100,000 EPS, legacy SIEMs produce ~40% false positives causing alert "
    "fatigue, and incident investigation averages 24–72 hours. Log evidence for forensics must maintain "
    "chain of custody, yet no major SIEM offers cryptographic tamper detection. Enterprise SIEMs "
    "(Splunk, Sentinel) cost $50K–$500K/yr — prohibitive for Indian agencies and SMBs."
)

doc.add_heading("2.3 Deliverables Mapping", level=2)
P(
    "The problem statement prescribes seven deliverables. CLIF addresses each:"
)
P(
    "The Web-Based Dashboard is a 14-page Next.js 14 SOC interface with real-time event feeds, keyword "
    "and semantic search, alert management, attack graph visualization (React Flow), evidence chain "
    "viewer, and system health monitoring — backed by 11 API routes. The Log Ingestion & Parsing Module "
    "uses Vector to collect from seven source types, normalize via VRL into the CLIF Common Schema, and "
    "sink to Redpanda; three Python consumers ingest into ClickHouse at 56,612 EPS with MITRE enrichment."
)
P(
    "The Database Storage Layer combines ClickHouse (ReplicatedMergeTree, 15-20x ZSTD, sub-100ms) with "
    "MinIO S3 cold archival and TTL lifecycle. The AI Correlation & Inference Engine uses a SQL-template "
    "rule engine (13+ MITRE techniques) and a four-agent DSPy pipeline with LanceDB RAG. Filtering & "
    "Search combines SQL filtering with semantic vector search over 494K+ embeddings. Automated "
    "Reporting supports CSV/JSON export plus LLM-generated MITRE-mapped incident reports. LLM Prompt "
    "Integration is realized through the semantic search interface and DSPy-optimized agent prompts."
)

doc.add_heading("2.4 Additional Requirements", level=2)
RP([("Tamper-Proof Chain of Custody: ",True,False),
    ("SHA-256 Merkle trees anchor log batches; proofs stored in MinIO with S3 Object Lock (WORM). "
     "Any modification cascades hash mismatches — blockchain-grade immutability critical for "
     "Section 65B admissibility.",False,False)])
RP([("Explainable AI: ",True,False),
    ("Each agent outputs confidence scores, rule-match rationale, investigation steps, and IOC "
     "validation results — compiled into human-readable narratives.",False,False)])
RP([("Cloud Deployment & SIEM/SOAR: ",True,False),
    ("18 Docker services, deployable on-premise or cloud. Horizontal scaling at every layer. "
     "SOAR via webhook notifications from the Reporter Agent.",False,False)])

doc.add_heading("2.5 Milestone Alignment", level=2)
P(
    "Phase 1 (Ingestion + Filtering): Vector aggregates 7 sources → CCS → Redpanda; consumers at "
    "56K EPS; dashboard keyword search + filters. "
    "Phase 2 (Secure Storage + AI Prototype): ClickHouse ReplicatedMergeTree RF=3; Merkle anchoring; "
    "LanceDB semantic search (494K+ embeddings). "
    "Phase 3 (GUI): 14-page dashboard — charts, alerts, React Flow attack graphs, health gauges. "
    "Phase 4 (Reporting + Benchmark): CSV/JSON export, LLM reports, 8-test benchmark — "
    "Grade A: 56,612 EPS, 0% data loss, 37.6 QPS, 61.4ms avg query."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3: LITERATURE REVIEW / EXISTING SOLUTIONS  (~2 pages)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Literature Review / Existing Solutions", level=1)

doc.add_heading("3.1 Industry SIEM Landscape", level=2)
P(
    "The global SIEM market is projected at $9.2B by 2027. Six platforms dominate with distinct "
    "trade-offs that informed CLIF's architectural decisions:"
)
T(["Platform","Architecture","Key Strength","Key Limitation"],[
    ("Splunk ES (Cisco, $28B)","Index-based distributed","SPL, 2000+ apps, Risk-Based Alerting","$50K–$500K/yr, heavy resources"),
    ("Elastic Security","Elasticsearch cluster","1000+ rules, ML anomaly, OSS core","JVM overhead, complex scaling"),
    ("Microsoft Sentinel","Azure cloud-native","300+ connectors, KQL, Fusion ML","Azure lock-in, per-GB pricing"),
    ("CrowdStrike Falcon","Index-free (LogScale)","150x faster search, Charlotte AI","Closed ecosystem, premium cost"),
    ("Google Chronicle","Google Cloud backend","Flat-rate, 12+ mo retention, VirusTotal","Limited customization, GCP only"),
    ("Wazuh (Open Source)","OSSEC + OpenSearch","Free, agent EDR, FIM, compliance","No ML/AI, scaling limits"),
])

doc.add_heading("3.2 Research Foundations", level=2)
P("CLIF's design draws from research in stream processing, semantic search, and evidence verification:")
for t,d in [
    ("Marz & Warren (2015), Big Data Systems — ","Lambda/Kappa architecture: raw preservation alongside real-time materialized views."),
    ("MITRE ATT&CK v14 (2024) — ","14 tactics, 200+ techniques for detection rule mapping and kill chain visualization."),
    ("Kent (2015), LANL Dataset — ","58-day dataset (17,684 hosts, 12,425 users) for realistic load simulations."),
    ("Reimann et al. (2023), Sentence Transformers for Logs — ","Embeddings outperform TF-IDF for log similarity; CLIF adopts all-MiniLM-L6-v2."),
    ("Khare et al. (2024), Merkle Trees for Forensics — ","SHA-256 trees for scalable tamper detection; adopted in CLIF's evidence service."),
    ("Khattab et al. (2023), DSPy — ","Declarative LLM pipeline compilation for reliable agent behavior; powers CLIF's agent chain."),
]:
    p = doc.add_paragraph(); p.add_run(t).bold = True; p.add_run(d)

doc.add_heading("3.3 Gaps in Existing Solutions", level=2)
P("Our 22-dimension competitive analysis identified five critical gaps:")
for t,d in [
    ("No cryptographic evidence verification — ","All SIEMs use append-only logs. None offer Merkle tree proof of integrity for forensic proceedings (Section 65B)."),
    ("AI/LLM integration is proprietary — ","Only Elastic and CrowdStrike have LLM assistants, both closed-source. No affordable SIEM offers semantic search."),
    ("Open-source SIEMs lack AI — ","Wazuh has no ML, no semantic search, struggles beyond ~15K EPS."),
    ("Enterprise cost is prohibitive — ","Splunk/Sentinel at $50K–$500K/yr excludes Indian agencies and SMBs."),
    ("No columnar + vector hybrid — ","No SIEM combines ClickHouse-grade analytics with vector similarity for semantic correlation."),
]:
    p = doc.add_paragraph(); p.add_run(t).bold = True; p.add_run(d)

P("These gaps define the opportunity CLIF addresses.",italic=True,sa=6)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4: PROPOSED SOLUTION AND TECHNICAL ARCHITECTURE  (~2 pages)
#
# Criteria stressed:
#   - Technical Maturity & Innovation (25-30%): detailed architecture choices
#   - Scalability & Deployment Readiness (20%): horizontal scaling, Docker, resource efficiency
#   - Technical Clarity (10-15%): precise specs, no hand-waving
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Proposed Solution and Technical Architecture", level=1)

doc.add_heading("4.1 Architecture Overview", level=2)
P(
    "CLIF is a three-plane system — Data, Intelligence, and Presentation — fully containerized as "
    "18 Docker services orchestrated via Docker Compose across three isolated networks (frontend, "
    "backend, storage). This section details the production infrastructure with exact specifications."
)

doc.add_heading("4.2 Data Plane", level=2)

# --- Collection ---
RP([("Log Collection (Vector + Tetragon): ",True,False),
    ("The Vector aggregator (v0.42.0, Rust-based) collects from seven source types — syslog (RFC 5424), "
     "HTTP/webhook, Docker container logs, file-based, journald, and kernel-level eBPF telemetry via "
     "Tetragon. Vector applies VRL (Vector Remap Language) transforms to normalize all sources into the "
     "CLIF Common Schema (CCS) — a standardized field set covering source/destination IPs, ports, event "
     "type/category/action, process metadata, network flow, and MITRE ATT&CK tactic/technique fields. "
     "This normalization at ingestion time enables cross-source correlation in downstream analytics.",False,False)])

# --- Streaming ---
RP([("Event Streaming (Redpanda): ",True,False),
    ("A 3-broker Redpanda cluster (v24.2.8, C++ native, Kafka-compatible) handles event streaming. "
     "Four topics (raw-logs, security-events, process-events, network-events) with 12 partitions each "
     "and replication factor 3 ensure every message is stored on all brokers. Redpanda was chosen over "
     "Apache Kafka for its C++ native architecture — no JVM overhead, 10x lower tail latency, "
     "zero-dependency operations (no ZooKeeper). Benchmarked at 205,000 produce EPS with LZ4 compression.",False,False)])

# --- Consumers ---
RP([("Ingestion Consumers (Python ×3): ",True,False),
    ("Three horizontally-scaled Python consumers form a single consumer group, each handling 4 of 12 "
     "partitions. The pipeline uses orjson (Rust-based JSON, 3-10x faster than stdlib), confluent-kafka "
     "(librdkafka C library), a Semaphore-based WriterPool with 4 flush threads, and columnar inserts via "
     "clickhouse-driver's native TCP protocol with LZ4 wire compression. Events are batched at 200,000 "
     "rows or 0.5s (whichever first) and flushed asynchronously — the main consumption loop never blocks "
     "on I/O. This pipelined design achieves 56,612 EPS sustained throughput with zero data loss "
     "across 2.5M test events. Scaling is linear: adding consumers increases throughput proportionally.",False,False)])

# --- Storage ---
RP([("Analytics Storage (ClickHouse): ",True,False),
    ("ClickHouse 24.8 runs as a 2-node replicated shard with a dedicated Keeper instance (ZooKeeper "
     "replacement). Four ReplicatedMergeTree tables (raw_logs, security_events, process_events, "
     "network_events) use ZSTD(3) compression achieving 15-20x ratios. Specialized indexes — token "
     "bloom filters on message text, MinMax on IPs and timestamps, set indexes on categories — deliver "
     "sub-100ms analytical query latency. Two materialized views (events_per_minute, severity_hourly) "
     "power real-time dashboard metrics. A TTL-driven tiered storage policy moves data automatically: "
     "hot (7 days, SSD) → warm (30 days) → cold (S3/MinIO, 90+ days).",False,False)])

# --- Cold + Evidence ---
RP([("Cold Storage & Evidence (MinIO + Merkle): ",True,False),
    ("A 3-node MinIO cluster (erasure-coded, S3-compatible) provides cold storage for ClickHouse TTL "
     "overflow and the evidence archive. The Merkle anchor service (merkle_anchor.py, 475 lines) "
     "constructs SHA-256 binary Merkle trees over log batches every 30 minutes, computing per-row "
     "hashes server-side via ClickHouse's SHA256() function. Merkle roots are persisted in an "
     "evidence_anchors table; full proof objects are uploaded to MinIO with S3 Object Lock (WORM). "
     "Verification re-computes the tree and compares roots — any tampered row produces a different hash.",False,False)])

doc.add_heading("4.3 Intelligence Plane", level=2)
RP([("Semantic Search (LanceDB): ",True,False),
    ("LanceDB runs as a FastAPI service embedding log text into 384-dimensional vectors using the "
     "all-MiniLM-L6-v2 sentence transformer (local inference, no API dependency). It continuously syncs "
     "from ClickHouse, maintaining 494,000+ indexed embeddings. Investigators query in natural language; "
     "the query is embedded into the same vector space and matched via approximate nearest neighbor "
     "(ANN) search, returning semantically similar events ranked by cosine similarity.",False,False)])

RP([("Multi-Agent Pipeline: ",True,False),
    ("Four agents collaborate in sequence, each with a specialized role. The Triage Agent processes "
     "100% of events through a SQL rule engine (13+ MITRE ATT&CK technique patterns: brute force T1110, "
     "lateral movement T1021, privilege escalation T1548, exfiltration T1041, DNS tunneling T1071.004, "
     "etc.) plus a DSPy classifier for ambiguous patterns, filtering down to <1% as signals with "
     "confidence >70%. The Hunter Agent receives signals and assembles context: entity expansion (±15min "
     "window around flagged entity in ClickHouse), semantic similarity search via LanceDB, and graph "
     "walks (User→Process→Network→IP). The Verifier Agent fact-checks enriched findings against "
     "external IOC databases (VirusTotal, AbuseIPDB) and validates source log integrity via Merkle "
     "proofs. The Reporter Agent generates structured Markdown incident reports with MITRE kill chain "
     "mapping, confidence scores, affected entities, and recommended remediation actions.",False,False)])

doc.add_heading("4.4 Presentation Plane", level=2)
P(
    "The SOC dashboard is a Next.js 14 application (TypeScript, Tailwind CSS, shadcn/ui) with 14 pages "
    "and 11 API routes. Key interfaces: Overview (real-time metrics from ClickHouse materialized views), "
    "Live Feed (event stream via HTTP polling), Search (keyword + semantic via LanceDB), Alerts "
    "(severity-classified security events), Attack Graph (React Flow visualization of MITRE kill chains), "
    "Evidence Chain (Merkle proof inspection), System Health (direct service health checks to all 18 "
    "containers), and Threat Intelligence (MITRE ATT&CK technique coverage). CSV and JSON export "
    "is available from the search interface."
)

doc.add_heading("4.5 Scalability and Deployment", level=2)
P(
    "Every layer scales horizontally without architectural changes. Redpanda brokers can be added to "
    "the cluster; partition count can increase from 12 to 48+ for higher parallelism. Consumer instances "
    "scale linearly — doubling consumers doubles throughput. ClickHouse supports adding shards for "
    "write throughput and replicas for read throughput. LanceDB replicas serve concurrent search load. "
    "The Docker Compose orchestration includes resource limits (CPU/memory) per service and supports "
    "production overrides via docker-compose.prod.yml. For large-scale deployment, Kubernetes Helm "
    "charts with Tetragon as a DaemonSet enable auto-scaling across cloud providers (AWS/Azure/GCP). "
    "Benchmark validation: 56,612 EPS on commodity hardware (5.1 GB RAM), 11,085 EPS per GB — "
    "demonstrating resource efficiency that scales predictably."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5: INNOVATION AND NOVELTY ELEMENTS  (~2 pages)
#
# Criteria stressed:
#   - Technical Maturity & Innovation (25-30%): what's new, what's first
#   - Impact & Sustainability (10%): why it matters long-term
#   - Market Viability (20%): competitive differentiation
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Innovation and Novelty Elements", level=1)

P(
    "CLIF introduces five technical innovations and four architectural innovations that collectively "
    "differentiate it from every existing SIEM — commercial or open-source."
)

doc.add_heading("5.1 Technical Innovations", level=2)

RP([("1. Cryptographic Evidence Chain (Blockchain-Grade Immutability): ",True,False),
    ("CLIF is the only SIEM that provides SHA-256 Merkle tree evidence anchoring with S3 Object Lock "
     "(WORM). No major SIEM — Splunk, Elastic, Sentinel, CrowdStrike, Chronicle, or Wazuh — offers "
     "cryptographic chain-of-custody verification. Splunk and Elastic rely on access controls; "
     "Sentinel uses append-only Azure Monitor logs. CLIF's approach makes any log modification "
     "mathematically provable, which is essential for forensic proceedings under Section 65B of the "
     "Indian Evidence Act and equivalent international standards. This is a genuine first in the "
     "SIEM domain — the Merkle service (475 lines) runs autonomously, anchoring evidence every 30 "
     "minutes with zero analyst intervention.",False,False)],sa=6)

RP([("2. AI-Powered Semantic Log Search: ",True,False),
    ("CLIF enables natural language queries against security logs. Using all-MiniLM-L6-v2 sentence "
     "transformers, 494,000+ log entries are embedded into 384-dimensional vectors for approximate "
     "nearest neighbor search. An investigator searching for \"suspicious login attempt from unknown "
     "IP after business hours\" receives semantically relevant results across all log tables — a "
     "capability that keyword search or regex fundamentally cannot provide. Only Elastic Security has "
     "begun exploring embedding-based search (via ELSER), but it remains proprietary and limited. "
     "CLIF implements this as an open, extensible FastAPI service with auto-sync from ClickHouse.",False,False)],sa=6)

RP([("3. Multi-Agent Autonomous Investigation: ",True,False),
    ("Unlike rule-only SIEMs, CLIF employs four specialized AI agents that collaborate to detect, "
     "investigate, verify, and report threats autonomously. Each agent has defined input/output "
     "contracts: Triage outputs Signal objects (confidence >70%), Hunter outputs EnrichedFindings "
     "(with ±15min contextual evidence), Verifier outputs ConfirmedIncident or FalsePositive "
     "(with IOC validation proof), and Reporter outputs structured Markdown reports with MITRE "
     "mapping. The pipeline uses DSPy for LLM prompt optimization — enabling measurable, reproducible "
     "agent behavior rather than fragile prompt engineering. This multi-agent approach reduces mean "
     "time to investigate (MTTI) from hours to minutes.",False,False)],sa=6)

RP([("4. Columnar + Vector Hybrid Query Engine: ",True,False),
    ("CLIF uniquely combines ClickHouse's columnar analytics (15-20x compression, sub-100ms queries "
     "on structured data) with LanceDB's vector similarity search (semantic queries on unstructured "
     "text). This hybrid enables both traditional SQL analytics (\"count events by severity in last "
     "hour\") and AI-powered semantic correlation (\"find similar attack patterns\") in a single "
     "platform. No existing SIEM offers this combination — they use either indexing (Elastic/Splunk), "
     "index-free log-structured approaches (CrowdStrike LogScale), or cloud-native backends (Sentinel/"
     "Chronicle). The hybrid architecture is validated by industry trends: CrowdStrike's move away "
     "from indexing confirms columnar superiority for log analytics.",False,False)],sa=6)

RP([("5. eBPF Kernel-Level Observability: ",True,False),
    ("Tetragon provides kernel-level process, network, and file telemetry via eBPF — capturing "
     "syscalls, process trees, and network connections without the performance overhead of user-space "
     "agents. This gives CLIF visibility that application-level logging cannot achieve: detecting "
     "rootkits that hide from userspace, monitoring container escape attempts at the kernel boundary, "
     "and capturing process lineage (parent-child trees) for forensic reconstruction.",False,False)],sa=6)

doc.add_heading("5.2 Architectural Innovations", level=2)
for t,d in [
    ("Redpanda over Kafka: ",
     "C++ native message broker eliminates JVM overhead — 10x lower tail latency, simpler operations "
     "(no ZooKeeper), same Kafka protocol compatibility. Validated by benchmarks: 205K produce EPS."),
    ("ClickHouse over Elasticsearch: ",
     "10-20x better compression on structured log data, faster aggregation queries, native S3 tiering "
     "— validated by CrowdStrike's own move away from indexing."),
    ("Pipelined Async Consumers: ",
     "Non-blocking flush pipeline with Semaphore-based WriterPool achieves 56,612 sustained EPS "
     "on commodity hardware. The main consumption loop never waits on database I/O."),
    ("Deduplication at Ingestion: ",
     "Vector VRL transforms deduplicate process and network events before pipeline entry, reducing "
     "storage costs and downstream noise for the AI agents."),
]:
    p = doc.add_paragraph()
    p.add_run(t).bold = True
    p.add_run(d)

doc.add_page_break()

# placeholder for remaining sections
P("[Sections 6–9 to follow.]", italic=True, sa=8)

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT} ({os.path.getsize(OUT)/1024:.1f} KB)")
